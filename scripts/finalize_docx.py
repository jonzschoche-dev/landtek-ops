#!/usr/bin/env python3
"""finalize_docx.py — the FINALIZER. Grounded markdown → a professional, MODERN-REPORT Word document.
Separates content (the grounded .md = source of truth) from presentation. $0, local (python-docx).

Design: a clean sans title block with a thin accent rule (no wasted title/TOC pages), section headers
each underlined by an understated accent rule, airy spacing, clean light-header tables, real clickable
hyperlinks (URL hidden), and a quiet footer (CONFIDENTIAL · Page X of Y). Front matter = everything
before the first '---'; the body follows.

  python3 scripts/finalize_docx.py in.md out.docx
"""
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

CONTENT_W = 6.3
ACCENT = "2F5496"      # accent blue — rules + links
H1COL = RGBColor(0x1F, 0x38, 0x64)
H2COL = RGBColor(0x44, 0x54, 0x6A)
TITLECOL = RGBColor(0x1F, 0x29, 0x37)
MUTED = RGBColor(0x6B, 0x72, 0x80)
NARROW = {"status", "exhibit", "date", "¶", "support"}


def _bottom_border(paragraph, color=ACCENT, sz=6, space=2):
    pPr = paragraph._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr"); b = OxmlElement("w:bottom")
    for k, v in (("val", "single"), ("sz", str(sz)), ("space", str(space)), ("color", color)):
        b.set(qn("w:" + k), v)
    bdr.append(b); pPr.append(bdr)


def _left_accent(paragraph, color=ACCENT, sz=18):
    pPr = paragraph._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr"); b = OxmlElement("w:left")
    for k, v in (("val", "single"), ("sz", str(sz)), ("space", "10"), ("color", color)):
        b.set(qn("w:" + k), v)
    bdr.append(b); pPr.append(bdr)


def _outline(paragraph, level):
    pPr = paragraph._p.get_or_add_pPr()
    o = OxmlElement("w:outlineLvl"); o.set(qn("w:val"), str(level)); pPr.append(o)


def _hyperlink(paragraph, url, text):
    r_id = paragraph.part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    h = OxmlElement("w:hyperlink"); h.set(qn("r:id"), r_id)
    run = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color"); c.set(qn("w:val"), ACCENT); rPr.append(c)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    run.append(rPr); t = OxmlElement("w:t"); t.text = text; run.append(t)
    h.append(run); paragraph._p.append(h)


def _field(paragraph, instr):
    r = paragraph.add_run()
    for typ in ("begin", "instr", "end"):
        if typ == "instr":
            e = OxmlElement("w:instrText"); e.set(qn("xml:space"), "preserve"); e.text = instr
        else:
            e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), typ)
        r._r.append(e)


def _inline(paragraph, text, size=None, color=None, italic=False):
    for part in re.split(r"(\*\*.+?\*\*|\*[^*\n]+?\*|`.+?`|\[[^\]]+\]\(https?://[^)\s]+\))", text):
        if not part:
            continue
        m = re.match(r"\[([^\]]+)\]\((https?://[^)\s]+)\)", part)
        if m:
            _hyperlink(paragraph, m.group(2), m.group(1)); continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2]); run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1]); run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1]); run.font.name = "Consolas"
        else:
            run = paragraph.add_run(part)
        if size:
            run.font.size = Pt(size)
        if color is not None:
            run.font.color.rgb = color
        if italic:
            run.italic = True


def _shade(cell, fill):
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def _widths(headers):
    n = len(headers)
    narrow = [i for i, h in enumerate(headers) if h.strip().lower() in NARROW]
    if not narrow or len(narrow) == n:
        return [CONTENT_W / n] * n
    nw = 0.9
    rest = (CONTENT_W - nw * len(narrow)) / (n - len(narrow))
    return [nw if i in narrow else rest for i in range(n)]


def _no_borders(t):
    tblPr = t._tbl.tblPr
    bd = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement("w:" + edge)
        for k, v in (("val", "single"), ("sz", "2"), ("space", "0"), ("color", "D5DBE2")):
            e.set(qn("w:" + k), v)
        bd.append(e)
    tblPr.append(bd)


def _table(doc, rows):
    headers = [c.strip() for c in rows[0]]
    n = len(headers)
    t = doc.add_table(rows=len(rows), cols=n)
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    t.autofit = False
    _no_borders(t)
    w = _widths(headers)
    for ri, row in enumerate(rows):
        for ci in range(n):
            cell = t.cell(ri, ci); cell.width = Inches(w[ci])
            cell.margins_top = None
            tcPr = cell._tc.get_or_add_tcPr()
            mar = OxmlElement("w:tcMar")
            for side, val in (("top", "60"), ("bottom", "60"), ("left", "100"), ("right", "100")):
                m = OxmlElement("w:" + side); m.set(qn("w:w"), val); m.set(qn("w:type"), "dxa"); mar.append(m)
            tcPr.append(mar)
            p = cell.paragraphs[0]; p.text = ""
            p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.1
            _inline(p, row[ci].strip() if ci < len(row) else "", size=8.5)
            if ri == 0:
                _shade(cell, "EAEFF5")
                for run in p.runs:
                    run.bold = True; run.font.color.rgb = H1COL
            elif ri % 2 == 0:
                _shade(cell, "F7F9FC")
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def build(md_path, out_path):
    doc = Document()
    n = doc.styles["Normal"]; n.font.name = "Calibri"; n.font.size = Pt(10.5)
    n.paragraph_format.line_spacing = 1.3; n.paragraph_format.space_after = Pt(6)
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Inches(1.0)
        sec.left_margin = sec.right_margin = Inches(1.1)
    # quiet footer
    fp = doc.sections[0].footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _bottom_border(fp, "E2E6EC", 4)  # actually a top hairline via paragraph border is awkward; keep simple
    fr = fp.add_run("CONFIDENTIAL  ·  Prepared by LandTek for counsel review      Page ")
    _field(fp, "PAGE"); fp.add_run(" of "); _field(fp, "NUMPAGES")
    for rr in fp.runs:
        rr.font.size = Pt(8); rr.font.color.rgb = MUTED

    lines = open(md_path).read().splitlines()
    div = next((i for i, l in enumerate(lines) if l.strip() == "---"), None)
    front = lines[:div] if div is not None else lines[:2]
    body = lines[div + 1:] if div is not None else lines[2:]

    # ── title block (no blank page, no TOC) ──
    title_done = False
    for raw in front:
        t = raw.strip()
        if not t:
            continue
        if t.startswith("# ") and not title_done:
            p = doc.add_paragraph(); r = p.add_run(t[2:]); r.bold = True
            r.font.size = Pt(23); r.font.color.rgb = TITLECOL
            p.paragraph_format.space_after = Pt(3)
            rule = doc.add_paragraph(); rule.paragraph_format.space_after = Pt(10)
            rr = rule.add_run(); rr.font.size = Pt(2); _bottom_border(rule, ACCENT, 10)
            title_done = True
        elif t.startswith("## "):
            p = doc.add_paragraph(); _inline(p, t[3:], size=12.5, color=RGBColor(0x5B, 0x64, 0x70))
            p.paragraph_format.space_after = Pt(12)
        else:
            p = doc.add_paragraph(); _inline(p, t, color=MUTED)

    # ── body ──
    rows = []
    def flush():
        if rows:
            _table(doc, rows); rows.clear()
    for raw in body:
        st = raw.strip()
        if st.startswith("|"):
            cells = [c.strip() for c in st.strip("|").split("|")]
            if set("".join(cells)) <= set("-: "):
                continue
            rows.append(cells); continue
        flush()
        if not st:
            continue
        if st.startswith("## "):
            p = doc.add_paragraph(); r = p.add_run(re.sub(r"\*\*", "", st[3:]))
            r.bold = True; r.font.size = Pt(13.5); r.font.color.rgb = H1COL
            p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(5)
            _bottom_border(p, ACCENT, 4); _outline(p, 0)
        elif st.startswith("### "):
            p = doc.add_paragraph(); r = p.add_run(re.sub(r"\*\*", "", st[4:]))
            r.bold = True; r.font.size = Pt(11.5); r.font.color.rgb = H2COL
            p.paragraph_format.space_before = Pt(11); p.paragraph_format.space_after = Pt(3); _outline(p, 1)
        elif st.startswith("# "):
            p = doc.add_paragraph(); r = p.add_run(re.sub(r"\*\*", "", st[2:]))
            r.bold = True; r.font.size = Pt(15); r.font.color.rgb = TITLECOL
        elif st.startswith(">"):
            p = doc.add_paragraph(); _inline(p, st.lstrip("> ").strip(), color=RGBColor(0x37, 0x41, 0x51))
            p.paragraph_format.left_indent = Inches(0.2); p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8); _left_accent(p, ACCENT, 22)
        elif re.match(r"^\*[^*].*[^*]\*$", st):     # *italic note*
            p = doc.add_paragraph(); _inline(p, st[1:-1], color=MUTED, italic=True)
            p.paragraph_format.space_after = Pt(10)
        elif st.startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet"); _inline(p, st[2:])
        elif re.match(r"\d+\.\s", st):
            p = doc.add_paragraph(style="List Number"); _inline(p, re.sub(r"^\d+\.\s", "", st))
        else:
            p = doc.add_paragraph(); _inline(p, st)
    flush()
    doc.save(out_path)


def main():
    if len(sys.argv) < 3:
        sys.exit("usage: finalize_docx.py in.md out.docx")
    build(sys.argv[1], sys.argv[2])
    print(f"[finalize] {sys.argv[2]}")


if __name__ == "__main__":
    main()
