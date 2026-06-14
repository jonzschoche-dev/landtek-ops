#!/usr/bin/env python3
"""extract_formats.py — drain the non-OCR extraction backlog: file types the OCR loop
can't read but that carry real text — .eml (email), .docx, .xlsx, .html. Creditless
(stdlib + python-docx/openpyxl if present). Fills documents.extracted_text + text_length.

The OCR loop handles scans/images; this handles structured documents. Physical-only vault
entries (no file) are correctly left empty.

  python3 extract_formats.py            # dry: list what it would extract
  python3 extract_formats.py --go       # extract + write
"""
import os
import re
import sys

import psycopg2
import psycopg2.extras

DSN = os.environ.get("PG_DSN", "postgresql://n8n:n8npassword@172.18.0.3:5432/n8n")


def _conn():
    c = psycopg2.connect(DSN); c.autocommit = True; return c


def _eml(path):
    import email
    from email import policy
    with open(path, "rb") as f:
        msg = email.message_from_bytes(f.read(), policy=policy.default)
    hdr = f"From: {msg.get('from','')}\nTo: {msg.get('to','')}\nSubject: {msg.get('subject','')}\nDate: {msg.get('date','')}\n\n"
    try:
        body = msg.get_body(preferencelist=("plain", "html"))
        text = body.get_content() if body else ""
    except Exception:
        text = ""
    if "<" in text and ">" in text:
        text = re.sub(r"<[^>]+>", " ", text)
    return hdr + text


def _docx(path):
    from docx import Document
    d = Document(path)
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for t in d.tables:
        for row in t.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts)


def _xlsx(path):
    import openpyxl
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    out = []
    for ws in wb.worksheets:
        out.append(f"# sheet: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                out.append(" | ".join(cells))
    return "\n".join(out)


def _html(path):
    with open(path, "r", errors="replace") as f:
        return re.sub(r"\s+", " ", re.sub(r"<[^>]+>", " ", f.read()))


def _pick(path, mime):
    low = (path or "").lower()
    m = (mime or "").lower()
    if low.endswith(".eml") or "rfc822" in m:
        return _eml
    if low.endswith(".docx") or "wordprocessing" in m:
        return _docx
    if low.endswith(".xlsx") or "spreadsheet" in m:
        return _xlsx
    if low.endswith((".html", ".htm")) or "text/html" in m:
        return _html
    return None


def run(go=False):
    c = _conn(); cur = c.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
    cur.execute("""SELECT id, file_path, mime_type, original_filename
        FROM documents
        WHERE length(coalesce(extracted_text,'')) < 50 AND coalesce(file_path,'') <> ''""")
    rows = cur.fetchall()
    done = skipped = failed = 0
    for r in rows:
        fn = _pick(r["file_path"], r["mime_type"])
        if not fn:
            skipped += 1
            continue
        if not os.path.exists(r["file_path"]):
            skipped += 1
            continue
        try:
            text = (fn(r["file_path"]) or "").strip()
        except Exception as e:
            failed += 1
            print(f"  ! doc {r['id']} ({r['mime_type']}): {type(e).__name__}: {str(e)[:80]}")
            continue
        if len(text) < 50:
            skipped += 1
            continue
        if go:
            cur.execute("UPDATE documents SET extracted_text=%s, text_length=%s WHERE id=%s",
                        (text[:200000], len(text), r["id"]))
        done += 1
        print(f"  {'extracted' if go else 'WOULD extract'} doc {r['id']} ({r['mime_type']}): {len(text)} chars")
    print(f"[extract_formats] {'WROTE' if go else 'DRY'} done={done} skipped={skipped} failed={failed}")
    cur.close(); c.close()


if __name__ == "__main__":
    run(go="--go" in sys.argv)
