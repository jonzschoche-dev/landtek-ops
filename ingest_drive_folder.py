#!/usr/bin/env python3
"""ingest_drive_folder.py — ingest ANY Google Drive folder into a named matter's corpus.

Generalized, repeatable sweep (supersedes the ingest_paracale_drive.py one-off). Point it at a
Drive folder ID and a case_file; it force-tags every item to that case_file (never auto-classifies,
so an Inocalla doc that mentions Keesey can't be misrouted to MWK), extracts text, embeds, and
logs to Qdrant + Postgres `documents`. Resumable via content_hash / drive_file_id dedup.

OCR ladder (the reason this exists): phone-photos of documents come back thin/empty from Document AI,
but Gemini vision transcribes them faithfully. So images (and scanned PDFs) go DocAI first, then fall
back to Gemini vision whenever the result is thin. --reocr-thin re-runs that vision pass over an
already-ingested batch to lift the blanks.

Provenance: extracted text only; doc_date left NULL (a Drive/export date is not the document date).
docx -> execution_status='draft'; pdf/image -> 'received'. Whole run carries one --source tag so the
batch is a single query to audit or roll back.

Run ON THE VPS (has creds, DocAI, DB). Source /root/landtek/.env first.

  python3 ingest_drive_folder.py --folder <ID> --case Paracale-001 --source <tag> --inventory
  python3 ingest_drive_folder.py --folder <ID> --case Paracale-001 --source <tag> --run
  python3 ingest_drive_folder.py --folder <ID> --case Paracale-001 --source <tag> --run --skip-photos
  python3 ingest_drive_folder.py --source <tag> --reocr-thin           # vision re-OCR the blanks of a batch
  python3 ingest_drive_folder.py --source <tag> --reocr-thin --go      # write them
"""
import os, sys, json, time, hashlib, base64, argparse, uuid, io
from datetime import datetime
import requests
import fitz  # PyMuPDF
import psycopg2
from google.oauth2 import service_account
from googleapiclient.discovery import build

GOOGLE_CREDS   = os.getenv("GOOGLE_APPLICATION_CREDENTIALS", "/root/landtek/google-creds.json")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")
GEMINI_KEY_FB  = os.getenv("GEMINI_API_KEY_FALLBACK", "")
QDRANT_URL = os.getenv("QDRANT_URL", ""); QDRANT_KEY = os.getenv("QDRANT_KEY", "")
QDRANT_COLL = "landtek_documents"
DOCAI_URL = "https://us-documentai.googleapis.com/v1/projects/287898704764/locations/us/processors/29ccddeea977ef1f:process"
PG_HOST = os.getenv("PGHOST", "172.18.0.3"); PG_USER = os.getenv("PGUSER", "n8n")
PG_PASSWORD = os.getenv("PGPASSWORD", "n8npassword"); PG_DB = os.getenv("PGDATABASE", "n8n")
EMBED_DIM = 768; CHUNK_SIZE = 400; CHUNK_OVERLAP = 50
MIN_TEXT_CHARS = 40          # below this -> catalog-only + eligible for vision re-OCR
EMBED_SLEEP = 1.0
VISION_MODELS = ["gemini-2.5-flash", "gemini-2.0-flash"]
VISION_PROMPT = ("Transcribe ALL text in this image faithfully and completely, preserving the "
                 "original wording, numbers, names and dates. This is a photographed legal/property/"
                 "mining document. Output only the transcription, no commentary. If the image has no "
                 "legible text, output exactly: [NO LEGIBLE TEXT].")
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
IMG_MIMES = ("image/jpeg", "image/jpg", "image/png")
SKIP_NAMES = (".DS_Store",)

def log(m): print(f"[{datetime.now().strftime('%H:%M:%S')}] {m}", flush=True)

# ---- auth / clients ---------------------------------------------------------
def drive():
    creds = service_account.Credentials.from_service_account_file(
        GOOGLE_CREDS, scopes=["https://www.googleapis.com/auth/drive"])
    return build("drive", "v3", credentials=creds, cache_discovery=False)

def docai_token():
    creds = service_account.Credentials.from_service_account_file(
        GOOGLE_CREDS, scopes=["https://www.googleapis.com/auth/cloud-platform"])
    import google.auth.transport.requests as gtr
    creds.refresh(gtr.Request()); return creds.token

def pg():
    return psycopg2.connect(host=PG_HOST, port=5432, dbname=PG_DB, user=PG_USER, password=PG_PASSWORD)

# ---- drive walk / download --------------------------------------------------
def walk(svc, root_id):
    out, queue, seen = [], [(root_id, "")], set()
    while queue:
        fid, top = queue.pop(0)
        if fid in seen: continue
        seen.add(fid)
        tok = None
        while True:
            r = svc.files().list(
                q=f"'{fid}' in parents and trashed=false",
                fields="nextPageToken, files(id,name,mimeType,size,modifiedTime)",
                pageSize=500, pageToken=tok,
                supportsAllDrives=True, includeItemsFromAllDrives=True).execute()
            for f in r.get("files", []):
                if f["name"] in SKIP_NAMES: continue
                if f["mimeType"].endswith(".folder"):
                    queue.append((f["id"], top or f["name"]))
                else:
                    f["_sub"] = top or "root"
                    out.append(f)
            tok = r.get("nextPageToken")
            if not tok: break
    return out

def download(svc, file_id):
    return svc.files().get_media(fileId=file_id, supportsAllDrives=True).execute()

# ---- OCR / extraction -------------------------------------------------------
def docai_ocr(file_bytes, mime):
    r = requests.post(DOCAI_URL,
        headers={"Authorization": f"Bearer {docai_token()}", "Content-Type": "application/json"},
        json={"rawDocument": {"content": base64.b64encode(file_bytes).decode(), "mimeType": mime}},
        timeout=180)
    r.raise_for_status()
    return r.json().get("document", {}).get("text", "")

def gemini_vision_ocr(img_bytes, mime):
    """Faithful transcription via Gemini vision. Economy ladder over models + fallback key."""
    b64 = base64.b64encode(img_bytes).decode()
    for key in [k for k in (GEMINI_API_KEY, GEMINI_KEY_FB) if k]:
        for model in VISION_MODELS:
            url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={key}"
            body = {"contents": [{"parts": [
                        {"text": VISION_PROMPT},
                        {"inline_data": {"mime_type": mime, "data": b64}}]}],
                    "generationConfig": {"temperature": 0.0, "maxOutputTokens": 4000}}
            try:
                r = requests.post(url, json=body, timeout=120)
                if r.status_code == 429:
                    time.sleep(20); continue
                if r.status_code >= 400:
                    continue
                txt = r.json()["candidates"][0]["content"]["parts"][0]["text"].strip()
                if txt == "[NO LEGIBLE TEXT]":
                    return ""
                return txt
            except Exception:
                continue
    return ""

def img_mime_of(name, mime):
    return "image/png" if (name.lower().endswith(".png") or mime == "image/png") else "image/jpeg"

def extract(file_bytes, mime, name):
    """Return (text, engine). DocAI/PyMuPDF first; Gemini vision rescues thin image/scan results."""
    try:
        if mime == "application/pdf" or file_bytes[:5] == b"%PDF-":
            with fitz.open(stream=file_bytes, filetype="pdf") as doc:
                text = "".join(p.get_text() for p in doc)
            if len(text) > 200 and sum(c.isalpha() for c in text) / max(len(text), 1) > 0.35:
                return text, "pymupdf"
            ocr = docai_ocr(file_bytes, "application/pdf")
            return (ocr, "docai") if len(ocr) >= MIN_TEXT_CHARS else (ocr, "docai-thin")
        if mime == DOCX_MIME or name.lower().endswith(".docx"):
            import docx
            d = docx.Document(io.BytesIO(file_bytes))
            parts = [p.text for p in d.paragraphs if p.text.strip()]
            for t in d.tables:
                for row in t.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells: parts.append(" | ".join(cells))
            return "\n".join(parts), "python-docx"
        if mime in IMG_MIMES or name.lower().endswith((".jpg", ".jpeg", ".png")):
            m = img_mime_of(name, mime)
            ocr = ""
            try: ocr = docai_ocr(file_bytes, m)
            except Exception as e: log(f"    docai err: {str(e)[:80]}")
            if len(ocr) >= MIN_TEXT_CHARS:
                return ocr, "docai"
            vis = gemini_vision_ocr(file_bytes, m)          # rescue thin/empty photos
            if len(vis) > len(ocr):
                return vis, "gemini-vision"
            return ocr, "docai-thin"
    except Exception as e:
        log(f"    extract error ({name}): {type(e).__name__}: {str(e)[:120]}")
        return "", "failed"
    return "", "unsupported"

# ---- embed / qdrant / postgres ---------------------------------------------
def chunk_text(text):
    words = text.split(); out, i = [], 0
    while i < len(words):
        out.append(" ".join(words[i:i+CHUNK_SIZE])); i += CHUNK_SIZE - CHUNK_OVERLAP
    return [c for c in out if len(c.strip()) > 50]

def embed(text):
    for attempt in range(4):
        r = requests.post(
            f"https://generativelanguage.googleapis.com/v1beta/models/gemini-embedding-001:embedContent?key={GEMINI_API_KEY}",
            json={"model": "models/gemini-embedding-001",
                  "content": {"parts": [{"text": text[:8000]}]}, "outputDimensionality": EMBED_DIM},
            timeout=60)
        if r.status_code == 429:
            time.sleep(15 * (attempt + 1)); continue
        r.raise_for_status()
        return r.json()["embedding"]["values"]
    raise RuntimeError("embed: exhausted retries on 429")

def qdrant_upsert(points):
    requests.put(f"{QDRANT_URL}/collections/{QDRANT_COLL}/points",
        headers={"api-key": QDRANT_KEY, "Content-Type": "application/json"},
        json={"points": points}, timeout=60).raise_for_status()

def embed_and_upsert(doc_id, content_hash, name, drive_id, case_file, source, text):
    chunks = chunk_text(text); pts = []
    for ci, ch in enumerate(chunks):
        try:
            vec = embed(ch); time.sleep(EMBED_SLEEP)
        except Exception as e:
            log(f"    embed chunk {ci} failed: {str(e)[:90]}"); continue
        pts.append({"id": str(uuid.uuid5(uuid.NAMESPACE_URL, f"{content_hash}:{ci}")),
                    "vector": vec, "payload": {"document_id": doc_id, "case_file": case_file,
                    "chunk_index": ci, "source_file": name, "drive_file_id": drive_id,
                    "ingest_source": source, "text": ch}})
    if pts: qdrant_upsert(pts)
    return len(pts)

def insert_doc(cur, *, case_file, source, name, mime, text, engine, content_hash, drive_id, sub, folder_root, chunk_count):
    is_docx = mime == DOCX_MIME or name.lower().endswith(".docx")
    exec_status = "draft" if is_docx else "received"
    doc_type = ("Working Draft" if is_docx else
                "Document (photographed)" if sub == "photos" else "Document")
    meta = {"source": "drive_folder_sweep", "drive_folder_root": folder_root, "sub": sub,
            "ocr_engine": engine, "original_name": name,
            "note": "doc_date NOT set (drive/export date != document date)"}
    cur.execute("""
        INSERT INTO documents
          (case_file, original_filename, smart_filename, file_name, document_title,
           mime_type, extracted_text, text_length, chunk_count, classification, document_type,
           execution_status, execution_metadata, content_hash, drive_file_id, drive_folder_id,
           ocr_used, ingest_source, ingest_status, status, confidence,
           first_seen_at, last_seen_at, created_at, processed_at)
        VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s::jsonb,%s,%s,%s,%s,%s,%s,%s,%s,
                NOW(),NOW(),NOW(),NOW())
        ON CONFLICT (content_hash) WHERE content_hash IS NOT NULL DO UPDATE SET
           last_seen_at=NOW(), duplicate_count=documents.duplicate_count+1
        RETURNING id, (xmax=0) AS inserted""",
        (case_file, name, name, name, name, mime, text[:200000], len(text), chunk_count,
         "Evidence", doc_type, exec_status, json.dumps(meta), content_hash, drive_id, folder_root,
         "docai" in engine or "vision" in engine, source, "ingested", "ingested",
         0.0 if not text.strip() else 0.6))
    return cur.fetchone()

# ---- modes ------------------------------------------------------------------
def doccount(case_file):
    conn = pg(); cur = conn.cursor()
    cur.execute("SELECT count(*) FROM documents WHERE case_file=%s", (case_file,))
    n = cur.fetchone()[0]; cur.close(); conn.close(); return n

def inventory(svc, folder_root, case_file):
    files = walk(svc, folder_root)
    conn = pg(); cur = conn.cursor()
    by_type = {}; total = 0; already = 0
    for f in files:
        ext = f["name"].rsplit(".", 1)[-1].lower() if "." in f["name"] else f["mimeType"]
        by_type[ext] = by_type.get(ext, 0) + 1
        total += int(f.get("size", 0) or 0)
        cur.execute("SELECT 1 FROM documents WHERE drive_file_id=%s LIMIT 1", (f["id"],))
        if cur.fetchone(): already += 1
    cur.close(); conn.close()
    log(f"Folder {folder_root} — {len(files)} files, {total/1e6:.1f} MB -> {case_file}")
    for k, v in sorted(by_type.items(), key=lambda x: -x[1]): log(f"   {v:>4}  .{k}")
    log(f"already ingested (by drive_file_id): {already}")
    log(f"current {case_file} doc count: {doccount(case_file)}")

def run(svc, folder_root, case_file, source, max_items=None, skip_photos=False):
    files = walk(svc, folder_root)
    if skip_photos: files = [f for f in files if f["_sub"] != "photos"]
    files.sort(key=lambda f: (f["_sub"] == "photos", f["name"]))
    if max_items: files = files[:max_items]
    stats = {"ingested": 0, "dup": 0, "empty": 0, "failed": 0, "vision": 0, "chunks": 0}
    for idx, f in enumerate(files, 1):
        name, mime, fid = f["name"], f["mimeType"], f["id"]
        log(f"[{idx}/{len(files)}] {f['_sub']}/{name} ({mime})")
        try: raw = download(svc, fid)
        except Exception as e: log(f"    download failed: {str(e)[:100]}"); stats["failed"] += 1; continue
        content_hash = hashlib.sha256(raw).hexdigest()
        conn = pg(); cur = conn.cursor()
        cur.execute("SELECT id FROM documents WHERE content_hash=%s OR drive_file_id=%s LIMIT 1", (content_hash, fid))
        row = cur.fetchone()
        if row:
            log(f"    already ingested as doc {row[0]} — skip"); stats["dup"] += 1
            cur.close(); conn.close(); continue
        text, engine = extract(raw, mime, name); text = (text or "").strip()
        if engine == "gemini-vision": stats["vision"] += 1
        if len(text) < MIN_TEXT_CHARS:
            log(f"    thin/empty ({len(text)} chars, {engine}) — catalog-only"); stats["empty"] += 1
        chunks = chunk_text(text)
        drow = insert_doc(cur, case_file=case_file, source=source, name=name, mime=mime, text=text,
                          engine=engine, content_hash=content_hash, drive_id=fid, sub=f["_sub"],
                          folder_root=folder_root, chunk_count=len(chunks))
        doc_id, inserted = drow; conn.commit()
        if not inserted:
            log(f"    hash-dup -> doc {doc_id}"); stats["dup"] += 1; cur.close(); conn.close(); continue
        n = 0
        try:
            n = embed_and_upsert(doc_id, content_hash, name, fid, case_file, source, text)
            cur.execute("UPDATE documents SET chunk_count=%s WHERE id=%s", (n, doc_id)); conn.commit()
            stats["chunks"] += n
        except Exception as e:
            log(f"    embed/upsert failed: {str(e)[:110]}")
        log(f"    -> doc {doc_id}: {len(text)} chars, {n} chunks [{engine}]")
        stats["ingested"] += 1
        cur.close(); conn.close()
    log(f"DONE. {json.dumps(stats)}")
    log(f"{case_file} doc count now: {doccount(case_file)}")

def reocr_thin(svc, source, go=False, thin=MIN_TEXT_CHARS):
    """Vision re-OCR the thin/blank image docs of a prior batch."""
    conn = pg(); cur = conn.cursor()
    cur.execute("""SELECT id, drive_file_id, original_filename, mime_type, text_length
                   FROM documents WHERE ingest_source=%s AND drive_file_id IS NOT NULL
                     AND (text_length < %s OR text_length IS NULL)
                     AND (mime_type LIKE 'image/%%' OR execution_metadata->>'ocr_engine' LIKE 'docai%%')
                   ORDER BY id""", (source, thin))
    rows = cur.fetchall()
    log(f"{len(rows)} thin/blank docs in batch '{source}'{' (DRY — no writes)' if not go else ''}")
    fixed = 0
    for did, fid, name, mime, tl in rows:
        try: raw = download(svc, fid)
        except Exception as e: log(f"  doc {did} {name}: download failed {str(e)[:60]}"); continue
        m = img_mime_of(name, mime or "")
        txt = gemini_vision_ocr(raw, m).strip()
        log(f"  doc {did} {name}: {tl or 0} -> {len(txt)} chars")
        if go and len(txt) >= thin:
            content_hash = hashlib.sha256(raw).hexdigest()
            cur.execute("SELECT case_file FROM documents WHERE id=%s", (did,)); cf = cur.fetchone()[0]
            cur.execute("""UPDATE documents SET extracted_text=%s, text_length=%s, ocr_used=true,
                           execution_metadata = jsonb_set(coalesce(execution_metadata,'{}'::jsonb),
                             '{ocr_engine}','\"gemini-vision-reocr\"'), last_seen_at=NOW()
                           WHERE id=%s""", (txt[:200000], len(txt), did))
            conn.commit()
            n = embed_and_upsert(did, content_hash, name, fid, cf, source, txt)
            cur.execute("UPDATE documents SET chunk_count=%s WHERE id=%s", (n, did)); conn.commit()
            fixed += 1
    cur.close(); conn.close()
    log(f"reocr-thin done. {'fixed '+str(fixed) if go else 'dry run'}")

# ---- main -------------------------------------------------------------------
if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--folder"); ap.add_argument("--case", default="Paracale-001")
    ap.add_argument("--source"); ap.add_argument("--inventory", action="store_true")
    ap.add_argument("--run", action="store_true"); ap.add_argument("--max", type=int)
    ap.add_argument("--skip-photos", action="store_true")
    ap.add_argument("--reocr-thin", action="store_true"); ap.add_argument("--go", action="store_true")
    a = ap.parse_args()
    for k in ("GEMINI_API_KEY", "QDRANT_URL", "QDRANT_KEY"):
        if not globals()[k]: log(f"FATAL: {k} not set (source /root/landtek/.env)"); sys.exit(1)
    svc = drive()
    if a.reocr_thin:
        if not a.source: log("FATAL: --reocr-thin needs --source"); sys.exit(1)
        reocr_thin(svc, a.source, go=a.go)
    elif a.inventory:
        if not a.folder: log("FATAL: --inventory needs --folder"); sys.exit(1)
        inventory(svc, a.folder, a.case)
    elif a.run:
        if not (a.folder and a.source): log("FATAL: --run needs --folder and --source"); sys.exit(1)
        run(svc, a.folder, a.case, a.source, max_items=a.max, skip_photos=a.skip_photos)
    else:
        ap.print_help()
