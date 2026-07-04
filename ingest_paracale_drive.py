#!/usr/bin/env python3
"""ingest_paracale_drive.py — ingest a specific Drive folder into the Paracale-001 (Allan/Shishir Inocalla) corpus.

Purpose-built one-off (NOT the auto-classifying bulk pipeline). Handles the Shishir Inocalla
Facebook/Messenger data export folder: PDFs + DOCX attachments (under files/) and 155 photographed
documents (under photos/). Every item is FORCE-tagged case_file='Paracale-001' — we never let a
classifier route an Inocalla doc to MWK-001 (some drafts mention Keesey only as a separation note,
which would trip auto-classification and risk cross-client contamination).

Stages are decoupled on purpose:
  * CORE ingest (this script): extract text (PyMuPDF / python-docx / Document AI OCR) -> chunk ->
    gemini-embedding-001 -> Qdrant + Postgres `documents`. No generative LLM needed, so it is robust
    to Gemini generateContent 429s / no-Ollama. This is what makes the docs searchable under Paracale-001.
  * ENRICH (separate later pass): document_type / smart_filename / analyst_memo via LLM. Not done here.

Provenance discipline: extracted text only. doc_date left NULL (FB export modifiedTime is the export
date, not the document date — never fabricate it). docx -> execution_status='draft'; pdf/image ->
'received' (received-via-Messenger attachment). Everything carries ingest_source='fb_export_shishir_20260704'
and execution_metadata so the whole batch is one query to audit or roll back.

Usage (run ON THE VPS):
  python3 ingest_paracale_drive.py --inventory          # count/type/dup, no spend, no writes
  python3 ingest_paracale_drive.py --run                # full ingest (extract+embed+insert)
  python3 ingest_paracale_drive.py --run --max 5        # cap total items (smoke test)
  python3 ingest_paracale_drive.py --run --skip-photos  # files/ only, defer the 155 photos
"""
import os, sys, json, time, hashlib, base64, argparse, uuid, io
from datetime import datetime
import requests
import fitz  # PyMuPDF
import psycopg2
from google.oauth2 import service_account
from googleapiclient.discovery import build

# ---- config (mirrors bulk_ingest_mwk.py so it reuses the same creds/DB/Qdrant) ----
FOLDER_ROOT   = os.getenv("PARACALE_DRIVE_ROOT", "1D7MgbVO7A5xUhtJfmPUQZjdjvVByxGNe")
CASE_FILE     = "Paracale-001"
INGEST_SOURCE = "fb_export_shishir_20260704"
GOOGLE_CREDS  = os.getenv("GOOGLE_APPLICATION_CREDENTIALS", "/root/landtek/google-creds.json")
GEMINI_API_KEY= os.getenv("GEMINI_API_KEY", "")
QDRANT_URL    = os.getenv("QDRANT_URL", "")
QDRANT_KEY    = os.getenv("QDRANT_KEY", "")
QDRANT_COLL   = "landtek_documents"
DOCAI_URL     = "https://us-documentai.googleapis.com/v1/projects/287898704764/locations/us/processors/29ccddeea977ef1f:process"
PG_HOST = os.getenv("PGHOST", "172.18.0.3"); PG_USER = os.getenv("PGUSER", "n8n")
PG_PASSWORD = os.getenv("PGPASSWORD", "n8npassword"); PG_DB = os.getenv("PGDATABASE", "n8n")
EMBED_DIM = 768; CHUNK_SIZE = 400; CHUNK_OVERLAP = 50
MIN_TEXT_CHARS = 40          # below this we still record the doc but flag ocr as thin/empty
EMBED_SLEEP = 1.0            # be gentle on the embedding endpoint

DOCX_MIME  = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
IMG_MIMES  = ("image/jpeg", "image/jpg", "image/png")
SKIP_NAMES = (".DS_Store",)

def log(m): print(f"[{datetime.now().strftime('%H:%M:%S')}] {m}", flush=True)

# ---- auth / clients ---------------------------------------------------------
def drive():
    creds = service_account.Credentials.from_service_account_file(
        GOOGLE_CREDS, scopes=["https://www.googleapis.com/auth/drive"])
    return build("drive", "v3", credentials=creds, cache_discovery=False)

def docai_token():
    creds = service_account.Credentials.from_service_account_file(
        GOOGLE_CREDS, scopes=["https://www.googleapis.com/auth/cloud-platform"])
    import google.auth.transport.requests as gtr
    creds.refresh(gtr.Request()); return creds.token

def pg():
    return psycopg2.connect(host=PG_HOST, port=5432, dbname=PG_DB, user=PG_USER, password=PG_PASSWORD)

# ---- drive walk / download --------------------------------------------------
def walk(svc, root_id):
    """BFS; return list of file dicts with _fbfolder = top-level subfolder name (files/photos/root)."""
    out, queue, seen = [], [(root_id, "")], set()
    while queue:
        fid, top = queue.pop(0)
        if fid in seen: continue
        seen.add(fid)
        tok = None
        while True:
            r = svc.files().list(
                q=f"'{fid}' in parents and trashed=false",
                fields="nextPageToken, files(id,name,mimeType,size,modifiedTime)",
                pageSize=500, pageToken=tok,
                supportsAllDrives=True, includeItemsFromAllDrives=True).execute()
            for f in r.get("files", []):
                if f["name"] in SKIP_NAMES: continue
                if f["mimeType"].endswith(".folder"):
                    queue.append((f["id"], top or f["name"]))
                else:
                    f["_fbfolder"] = top or "root"
                    out.append(f)
            tok = r.get("nextPageToken")
            if not tok: break
    return out

def download(svc, file_id):
    return svc.files().get_media(fileId=file_id, supportsAllDrives=True).execute()

# ---- extraction -------------------------------------------------------------
def docai_ocr(file_bytes, mime):
    b64 = base64.b64encode(file_bytes).decode()
    r = requests.post(DOCAI_URL,
        headers={"Authorization": f"Bearer {docai_token()}", "Content-Type": "application/json"},
        json={"rawDocument": {"content": b64, "mimeType": mime}}, timeout=180)
    r.raise_for_status()
    return r.json().get("document", {}).get("text", "")

def extract(file_bytes, mime, name):
    """Return (text, engine). Never raises — worst case returns ('', 'failed')."""
    try:
        if mime == "application/pdf" or file_bytes[:5] == b"%PDF-":
            with fitz.open(stream=file_bytes, filetype="pdf") as doc:
                text = "".join(p.get_text() for p in doc)
            if len(text) > 200 and sum(c.isalpha() for c in text) / max(len(text), 1) > 0.35:
                return text, "pymupdf"
            # scanned/near-empty PDF -> OCR
            return docai_ocr(file_bytes, "application/pdf"), "docai"
        if mime == DOCX_MIME or name.lower().endswith(".docx"):
            import docx
            d = docx.Document(io.BytesIO(file_bytes))
            parts = [p.text for p in d.paragraphs if p.text.strip()]
            for t in d.tables:
                for row in t.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells: parts.append(" | ".join(cells))
            return "\n".join(parts), "python-docx"
        if mime in IMG_MIMES or name.lower().endswith((".jpg", ".jpeg", ".png")):
            m = "image/png" if name.lower().endswith(".png") or mime == "image/png" else "image/jpeg"
            return docai_ocr(file_bytes, m), "docai"
    except Exception as e:
        log(f"    extract error ({name}): {type(e).__name__}: {str(e)[:140]}")
        return "", "failed"
    return "", "unsupported"

# ---- embed / qdrant / postgres ---------------------------------------------
def chunk_text(text):
    words = text.split(); out, i = [], 0
    while i < len(words):
        out.append(" ".join(words[i:i+CHUNK_SIZE])); i += CHUNK_SIZE - CHUNK_OVERLAP
    return [c for c in out if len(c.strip()) > 50]

def embed(text):
    for attempt in range(4):
        r = requests.post(
            f"https://generativelanguage.googleapis.com/v1beta/models/gemini-embedding-001:embedContent?key={GEMINI_API_KEY}",
            json={"model": "models/gemini-embedding-001",
                  "content": {"parts": [{"text": text[:8000]}]},
                  "outputDimensionality": EMBED_DIM}, timeout=60)
        if r.status_code == 429:
            time.sleep(15 * (attempt + 1)); continue
        r.raise_for_status()
        return r.json()["embedding"]["values"]
    raise RuntimeError("embed: exhausted retries on 429")

def qdrant_upsert(points):
    requests.put(f"{QDRANT_URL}/collections/{QDRANT_COLL}/points",
        headers={"api-key": QDRANT_KEY, "Content-Type": "application/json"},
        json={"points": points}, timeout=60).raise_for_status()

def existing_by_hash_or_drive(cur, content_hash, drive_id):
    cur.execute("SELECT id FROM documents WHERE content_hash=%s OR drive_file_id=%s LIMIT 1",
                (content_hash, drive_id))
    row = cur.fetchone(); return row[0] if row else None

def insert_doc(cur, *, name, mime, text, engine, content_hash, drive_id, fbfolder, chunk_count):
    is_docx = mime == DOCX_MIME or name.lower().endswith(".docx")
    exec_status = "draft" if is_docx else "received"
    doc_type = ("Working Draft" if is_docx else
                "Document (photographed)" if fbfolder == "photos" else "Document")
    meta = {"source": "facebook_export", "fb_export": "shishirinocalla_10163918383770056",
            "fb_folder": fbfolder, "ocr_engine": engine, "original_name": name,
            "note": "FB/Messenger attachment; doc_date NOT set (export date != document date)"}
    cur.execute("""
        INSERT INTO documents
          (case_file, original_filename, smart_filename, file_name, document_title,
           mime_type, extracted_text, text_length, chunk_count, classification, document_type,
           execution_status, execution_metadata, content_hash, drive_file_id, drive_folder_id,
           ocr_used, ingest_source, ingest_status, status, confidence,
           first_seen_at, last_seen_at, created_at, processed_at)
        VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s::jsonb,%s,%s,%s,%s,%s,%s,%s,%s,
                NOW(),NOW(),NOW(),NOW())
        ON CONFLICT (content_hash) WHERE content_hash IS NOT NULL DO UPDATE SET
           last_seen_at=NOW(), duplicate_count=documents.duplicate_count+1
        RETURNING id, (xmax=0) AS inserted""",
        (CASE_FILE, name, name, name, name,
         mime, text[:200000], len(text), chunk_count, "Evidence", doc_type,
         exec_status, json.dumps(meta), content_hash, drive_id, FOLDER_ROOT,
         engine == "docai", INGEST_SOURCE, "ingested", "ingested",
         0.0 if not text.strip() else 0.6))
    return cur.fetchone()

# ---- inventory --------------------------------------------------------------
def inventory(svc):
    files = walk(svc, FOLDER_ROOT)
    conn = pg(); cur = conn.cursor()
    by_type = {}; total = 0; already = 0
    for f in files:
        ext = f["name"].rsplit(".", 1)[-1].lower() if "." in f["name"] else f["mimeType"]
        by_type[ext] = by_type.get(ext, 0) + 1
        total += int(f.get("size", 0) or 0)
        cur.execute("SELECT 1 FROM documents WHERE drive_file_id=%s LIMIT 1", (f["id"],))
        if cur.fetchone(): already += 1
    cur.close(); conn.close()
    log(f"Folder {FOLDER_ROOT} — {len(files)} files, {total/1e6:.1f} MB")
    for k, v in sorted(by_type.items(), key=lambda x: -x[1]):
        log(f"   {v:>4}  .{k}")
    log(f"already ingested (by drive_file_id): {already}")
    log(f"current Paracale-001 doc count: {doccount()}")
    return files

def doccount():
    conn = pg(); cur = conn.cursor()
    cur.execute("SELECT count(*) FROM documents WHERE case_file=%s", (CASE_FILE,))
    n = cur.fetchone()[0]; cur.close(); conn.close(); return n

# ---- run --------------------------------------------------------------------
def run(svc, max_items=None, skip_photos=False):
    files = walk(svc, FOLDER_ROOT)
    if skip_photos:
        files = [f for f in files if f["_fbfolder"] != "photos"]
    # files/ first (higher-value native docs), then photos
    files.sort(key=lambda f: (f["_fbfolder"] == "photos", f["name"]))
    if max_items: files = files[:max_items]
    stats = {"ingested": 0, "dup": 0, "empty": 0, "failed": 0, "chunks": 0}
    for idx, f in enumerate(files, 1):
        name, mime, fid = f["name"], f["mimeType"], f["id"]
        log(f"[{idx}/{len(files)}] {f['_fbfolder']}/{name} ({mime})")
        try:
            raw = download(svc, fid)
        except Exception as e:
            log(f"    download failed: {str(e)[:120]}"); stats["failed"] += 1; continue
        content_hash = hashlib.sha256(raw).hexdigest()
        conn = pg(); cur = conn.cursor()
        dup = existing_by_hash_or_drive(cur, content_hash, fid)
        if dup:
            log(f"    already ingested as doc {dup} — skip"); stats["dup"] += 1
            cur.close(); conn.close(); continue
        text, engine = extract(raw, mime, name)
        text = (text or "").strip()
        if len(text) < MIN_TEXT_CHARS:
            log(f"    thin/empty text ({len(text)} chars, engine={engine}) — recording as catalog-only")
            stats["empty"] += 1
        chunks = chunk_text(text)
        # insert the document row first (so it is recorded even if embedding later fails)
        row = insert_doc(cur, name=name, mime=mime, text=text, engine=engine,
                         content_hash=content_hash, drive_id=fid, fbfolder=f["_fbfolder"],
                         chunk_count=len(chunks))
        doc_id, inserted = row
        conn.commit()
        if not inserted:
            log(f"    hash-dup on insert -> doc {doc_id}"); stats["dup"] += 1
            cur.close(); conn.close(); continue
        # embed + upsert chunks
        pts = []
        for ci, ch in enumerate(chunks):
            try:
                vec = embed(ch); time.sleep(EMBED_SLEEP)
            except Exception as e:
                log(f"    embed chunk {ci} failed: {str(e)[:100]}"); continue
            pid = str(uuid.uuid5(uuid.NAMESPACE_URL, f"{content_hash}:{ci}"))
            pts.append({"id": pid, "vector": vec, "payload": {
                "document_id": doc_id, "case_file": CASE_FILE, "chunk_index": ci,
                "source_file": name, "drive_file_id": fid, "ingest_source": INGEST_SOURCE,
                "text": ch}})
        if pts:
            try:
                qdrant_upsert(pts)
                cur.execute("UPDATE documents SET chunk_count=%s WHERE id=%s", (len(pts), doc_id))
                conn.commit()
                stats["chunks"] += len(pts)
            except Exception as e:
                log(f"    qdrant upsert failed: {str(e)[:120]}")
        log(f"    -> doc {doc_id}: {len(text)} chars, {len(pts)} chunks embedded [{engine}]")
        stats["ingested"] += 1
        cur.close(); conn.close()
    log(f"DONE. {json.dumps(stats)}")
    log(f"Paracale-001 doc count now: {doccount()}")

# ---- main -------------------------------------------------------------------
if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--inventory", action="store_true")
    ap.add_argument("--run", action="store_true")
    ap.add_argument("--max", type=int, default=None)
    ap.add_argument("--skip-photos", action="store_true")
    a = ap.parse_args()
    for k in ("GEMINI_API_KEY", "QDRANT_URL", "QDRANT_KEY"):
        if not globals()[k]:
            log(f"FATAL: {k} not set (source /root/landtek/.env first)"); sys.exit(1)
    svc = drive()
    if a.inventory:
        inventory(svc)
    elif a.run:
        run(svc, max_items=a.max, skip_photos=a.skip_photos)
    else:
        ap.print_help()
