#!/usr/bin/env python3
"""Extract text from an uploaded file (PDF / DOCX / images).

Called by the n8n 'Extract File Text' Code node.

Input:  stdin JSON {base64_data, original_filename, mime_type?}
Output: stdout JSON {extracted_text, char_count, status, local_path, mime_type}

Strategies by extension:
  .pdf  -> fitz (PyMuPDF) text layer. If <200 chars, fall back to Gemini Vision.
  .docx -> python-docx paragraph extraction.
  .txt  -> read as utf-8.
  .png/.jpg/.jpeg -> Gemini Vision.
  others -> 'unsupported' status, no text.
"""
import base64
import json
import os
import sys
import re
import tempfile
from pathlib import Path

UPLOADS_DIR = Path("/root/landtek/uploads")
UPLOADS_DIR.mkdir(parents=True, exist_ok=True)


def extract_pdf(path: Path) -> str:
    """Extract via text-layer. Caller handles fallback if <200 chars."""
    import fitz
    doc = fitz.open(str(path))
    text = "\n".join(p.get_text() for p in doc)
    doc.close()
    return text.strip()


def extract_pdf_with_ocr_fallback(path: Path):
    """PDF extraction priority:
       1. fitz text-layer (free, native PDFs)
       2. Gemini 2.5 Flash PDF-native (cheap, layout-aware, ~10x cheaper than DocAI)
       3. Document AI (last resort if Gemini fails)

    Returns (text, status).
    """
    text = extract_pdf(path)
    if len(text) >= 200:
        return text, "ok_text_layer"
    # Primary fallback: Gemini PDF-native
    try:
        gem_text = ocr_via_gemini_pdf(path)
        if gem_text and len(gem_text) > len(text):
            return gem_text, "ok_gemini_pdf"
    except Exception as e:
        gem_err = f"{type(e).__name__}: {str(e)[:200]}"
    else:
        gem_err = "no_text_returned"
    # Secondary fallback: Document AI
    try:
        ocr_text = ocr_via_document_ai(path)
        if ocr_text and len(ocr_text) > len(text):
            return ocr_text, f"ok_document_ai (gemini_failed: {gem_err})"
    except Exception as e:
        return text, f"both_failed (gemini: {gem_err}; docai: {type(e).__name__}: {e})"
    return text, f"fallback_empty (gemini: {gem_err})"


def ocr_via_gemini_pdf(pdf_path: Path) -> str:
    """Gemini 2.5 Flash PDF-native extraction via Files API.

    Uploads the PDF to Gemini Files API, lets it process, then asks for the
    full text. Far cheaper than Document AI and preserves layout/structure.
    Reuses the timeout-wrapped pattern from heightened_ocr/.
    """
    from dotenv import load_dotenv
    load_dotenv("/root/landtek/.env")
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        return ""
    import google.generativeai as genai
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel("gemini-2.5-flash")
    # Upload the PDF
    uploaded = genai.upload_file(str(pdf_path), mime_type="application/pdf")
    # Wait for it to finish processing (up to 60s for large files)
    import time as _time
    for _ in range(60):
        f = genai.get_file(uploaded.name)
        if f.state.name == "ACTIVE":
            break
        if f.state.name == "FAILED":
            raise RuntimeError(f"gemini file processing failed: {f.state}")
        _time.sleep(1)
    else:
        raise TimeoutError("gemini file processing did not become ACTIVE in 60s")
    # Generate
    resp = model.generate_content([
        "Extract ALL text from this PDF document, preserving structure: headings, "
        "paragraphs, tables, page numbers, signatures, and any annotations. "
        "Return text only — no commentary, no markdown fences. "
        "For each page boundary, insert '\n--- Page N ---\n' as a separator.",
        uploaded,
    ], generation_config={"temperature": 0.0, "max_output_tokens": 65536})
    text = resp.text.strip() if resp and resp.text else ""
    # Cleanup uploaded file
    try:
        genai.delete_file(uploaded.name)
    except Exception:
        pass
    return text


def ocr_via_document_ai(pdf_path: Path) -> str:
    """Google Document AI OCR with chunking for PDFs >30 pages.

    Document AI synchronous endpoint has a 30-page cap. We split the PDF
    into 30-page chunks via fitz, process each, and concatenate.
    """
    from dotenv import load_dotenv
    import fitz
    load_dotenv("/root/landtek/.env")
    project = os.environ.get("DOCAI_PROJECT", "landtek")
    location = os.environ.get("DOCAI_LOCATION", "us")
    processor_id = os.environ.get("DOCAI_PROCESSOR", "29ccddeea977ef1f")
    if not (project and processor_id):
        return ""
    from google.cloud import documentai
    opts = {"api_endpoint": f"{location}-documentai.googleapis.com"} if location != "us" else None
    client = documentai.DocumentProcessorServiceClient(client_options=opts) if opts else documentai.DocumentProcessorServiceClient()
    name = client.processor_path(project, location, processor_id)

    src = fitz.open(str(pdf_path))
    total_pages = src.page_count
    if total_pages <= 15:
        # Single-shot
        raw = pdf_path.read_bytes()
        doc_input = documentai.RawDocument(content=raw, mime_type="application/pdf")
        request = documentai.ProcessRequest(name=name, raw_document=doc_input)
        result = client.process_document(request=request)
        src.close()
        return result.document.text.strip() if result and result.document else ""

    # Chunk into 30-page blocks
    chunks_text = []
    for start in range(0, total_pages, 15):
        end = min(start + 15, total_pages)
        chunk = fitz.open()
        chunk.insert_pdf(src, from_page=start, to_page=end - 1)
        chunk_bytes = chunk.tobytes()
        chunk.close()
        try:
            doc_input = documentai.RawDocument(content=chunk_bytes, mime_type="application/pdf")
            request = documentai.ProcessRequest(name=name, raw_document=doc_input)
            result = client.process_document(request=request)
            page_label = f"--- Pages {start+1}-{end} ---"
            text = result.document.text.strip() if result and result.document else ""
            chunks_text.append(f"{page_label}\n{text}")
        except Exception as e:
            chunks_text.append(f"--- Pages {start+1}-{end}: chunk error {type(e).__name__}: {e} ---")
    src.close()
    return "\n\n".join(chunks_text)


def extract_image_gemini_from_bytes(data: bytes, mime: str) -> str:
    import google.generativeai as genai
    from dotenv import load_dotenv
    load_dotenv("/root/landtek/.env")
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        return ""
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel("gemini-2.5-flash")
    img = {"mime_type": mime, "data": data}
    resp = model.generate_content([
        "Extract ALL visible text from this document image, preserving structure (headings, paragraphs, tables). Return text only, no commentary.",
        img,
    ])
    return resp.text.strip() if resp and resp.text else ""


def extract_docx(path: Path) -> str:
    import docx
    d = docx.Document(str(path))
    parts = [p.text for p in d.paragraphs if p.text and p.text.strip()]
    # Also pull text from tables
    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text and cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n".join(parts).strip()


def extract_image_gemini(path: Path) -> str:
    import google.generativeai as genai
    from dotenv import load_dotenv
    load_dotenv("/root/landtek/.env")
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        return ""
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel("gemini-2.5-flash")
    with open(path, "rb") as f:
        data = f.read()
    img = {"mime_type": "image/jpeg" if str(path).lower().endswith(("jpg", "jpeg")) else "image/png",
           "data": data}
    resp = model.generate_content([
        "Extract ALL visible text from this image, preserving structure. Return text only, no commentary.",
        img,
    ])
    return resp.text.strip() if resp and resp.text else ""


def main():
    raw = sys.stdin.read().strip()
    if not raw:
        print(json.dumps({"status": "error", "error": "empty stdin"}))
        return
    try:
        data = json.loads(raw)
    except Exception as e:
        print(json.dumps({"status": "error", "error": f"json parse: {e}"}))
        return

    b64 = data.get("base64_data") or ""
    original_filename = data.get("original_filename") or "uploaded_file"
    declared_mime = data.get("mime_type") or ""

    safe_name = re.sub(r"[^A-Za-z0-9._-]", "_", original_filename)
    local_path = UPLOADS_DIR / f"{os.getpid()}_{safe_name}"
    try:
        binary = base64.b64decode(b64)
    except Exception as e:
        print(json.dumps({"status": "error", "error": f"b64 decode: {e}"}))
        return
    local_path.write_bytes(binary)

    ext = local_path.suffix.lower()
    text = ""
    status = "ok"
    try:
        if ext == ".pdf":
            text, status = extract_pdf_with_ocr_fallback(local_path)
        elif ext == ".docx":
            text = extract_docx(local_path)
        elif ext in (".txt", ".md", ".csv"):
            text = local_path.read_text(encoding="utf-8", errors="replace")
        elif ext in (".png", ".jpg", ".jpeg", ".webp"):
            text = extract_image_gemini(local_path)
            status = "ok_gemini"
        else:
            status = f"unsupported_extension:{ext}"
    except Exception as e:
        status = f"extract_error:{type(e).__name__}: {e}"
        text = ""

    print(json.dumps({
        "extracted_text": text,
        "char_count": len(text),
        "status": status,
        "local_path": str(local_path),
        "mime_type": declared_mime or {
            ".pdf": "application/pdf",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".txt": "text/plain",
            ".png": "image/png",
            ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg",
        }.get(ext, "application/octet-stream"),
    }))


if __name__ == "__main__":
    main()
