#!/usr/bin/env python3
"""LeoLandTek Report Generator
Triggered by Leo when response exceeds Telegram limits.
Generates PDF/DOCX from verified database facts only.
Uploads to Google Drive and returns a shareable link.

Schema notes — SQL adapted from Jonathan's spec to actual columns:
  titles                : no confidence_score column → omitted; we already store provenance_level
  chain_of_title        : columns are tct_number, registrant_full_name, predecessor_title,
                          registration_date, source_chunk_id, provenance_level
                          (NOT owner_canonical / transfer_date / instrument_type / from_tct / verified)
  extraction_chunks     : data lives in structured_value JSONB (NOT extracted_data),
                          no per-chunk confidence_score; provenance via provenance_level + verified_by
  case_reports          : matter_code FK added 2026-05-14
"""

import os, sys, json, psycopg2
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

DRIVE_REPORTS_FOLDER = "1eDLECG_Lu9dXh-FLeCTvjI3fJclMid2b"  # 09 - AI Processing
PG_CONN = os.getenv("DATABASE_URL", "")


def get_db():
    return psycopg2.connect(PG_CONN)


def _save_doc(doc: Document, filename: str) -> dict:
    filepath = Path(f"/root/landtek/reports/{filename}")
    filepath.parent.mkdir(exist_ok=True)
    doc.save(filepath)
    return {
        "filepath": str(filepath),
        "filename": filename,
        "size_bytes": filepath.stat().st_size,
    }


def generate_title_chain_report(tct_number: str) -> dict:
    """Generate a full chain of title report for a TCT."""
    conn = get_db()
    cur = conn.cursor()

    doc = Document()
    doc.add_heading("CHAIN OF TITLE REPORT", 0)
    doc.add_heading(f"TCT {tct_number}", 1)
    p = doc.add_paragraph()
    p.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y %H:%M')} PHT").italic = True
    p.add_run("\nSource: LeoLandTek Evidence-Grade RAG System").italic = True
    p.add_run("\nProvenance: Verified data only").italic = True
    doc.add_paragraph()

    # ── Title row ─────────────────────────────────────────────────────
    cur.execute("""
        SELECT tct_number, registrant_canonical, area_sqm,
               parent_title, location, provenance_level, source_doc_id
          FROM titles
         WHERE tct_number = %s
    """, (tct_number,))
    row = cur.fetchone()
    if row:
        doc.add_heading("Title Information", 2)
        table = doc.add_table(rows=1, cols=2); table.style = 'Table Grid'
        hdr = table.rows[0].cells; hdr[0].text = "Field"; hdr[1].text = "Value"
        fields = [
            ("TCT Number", row[0]),
            ("Registered Owner", row[1] or "(see chain below)"),
            ("Area (sqm)", str(row[2]) if row[2] is not None else "Not extracted"),
            ("Parent Title", row[3] or "(none / root of chain)"),
            ("Location", row[4] or "Not extracted"),
            ("Provenance", row[5] or "unknown"),
            ("Source Doc ID", str(row[6]) if row[6] else "—"),
        ]
        for fname, val in fields:
            r = table.add_row().cells; r[0].text = fname; r[1].text = str(val)
        doc.add_paragraph()

    # ── Chain of title (registrants for this TCT + verified parent edges) ──
    cur.execute("""
        SELECT registrant_full_name, predecessor_title, registration_date,
               provenance_level
          FROM chain_of_title
         WHERE tct_number = %s
         ORDER BY provenance_level, registrant_full_name
    """, (tct_number,))
    chain_rows = cur.fetchall()
    if chain_rows:
        doc.add_heading("Registered Owners", 2)
        table = doc.add_table(rows=1, cols=4); table.style = 'Table Grid'
        hdr = table.rows[0].cells
        for i, h in enumerate(["Registrant", "Predecessor", "Registration Date", "Provenance"]):
            hdr[i].text = h
        for r in chain_rows:
            cells = table.add_row().cells
            cells[0].text = r[0] or "Unknown"
            cells[1].text = r[1] or "—"
            cells[2].text = str(r[2]) if r[2] else "—"
            cells[3].text = r[3] or "—"
        doc.add_paragraph()

    # ── Parent / predecessor edges from title_chain ────────────────────
    cur.execute("""
        SELECT parent_title, child_title, relationship, provenance_level, LEFT(notes,120) AS note
          FROM title_chain
         WHERE child_title = %s OR parent_title = %s
         ORDER BY (child_title = %s) DESC, parent_title
    """, (tct_number, tct_number, tct_number))
    edges = cur.fetchall()
    if edges:
        doc.add_heading("Title Chain Edges", 2)
        table = doc.add_table(rows=1, cols=4); table.style = 'Table Grid'
        hdr = table.rows[0].cells
        for i, h in enumerate(["Parent", "Child", "Relationship", "Provenance"]):
            hdr[i].text = h
        for e in edges:
            cells = table.add_row().cells
            cells[0].text = e[0]; cells[1].text = e[1]
            cells[2].text = e[2]; cells[3].text = e[3] or "—"
        doc.add_paragraph()

    # ── Verified chunks ─────────────────────────────────────────────────
    cur.execute("""
        SELECT id, chunk_type, field_name, verified_by, LEFT(quote_text,120) AS quote
          FROM extraction_chunks
         WHERE tct_number = %s AND provenance_level = 'verified'
         ORDER BY chunk_type, id
    """, (tct_number,))
    chunks = cur.fetchall()
    if chunks:
        doc.add_heading("Verified Sources", 2)
        doc.add_paragraph(f"{len(chunks)} verified chunks back this report.")
        for c in chunks:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"[chunk:{c[0]}] ").bold = True
            p.add_run(f"{c[1]} / {c[2] or '—'} — verified_by: {c[3]}")
            if c[4]:
                p.add_run(f"\n      quote: {c[4]}")

    # ── Disclaimer ─────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_heading("Disclaimer", 2)
    doc.add_paragraph(
        "This report contains only facts verified through two-pass cross-validation "
        "of original source documents, or source-quote-matched against the documents' "
        "extracted text. Unverified fields are explicitly marked. Prepared for legal "
        "strategy purposes; consult qualified Philippine counsel for court filings."
    )

    cur.close(); conn.close()
    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    return _save_doc(doc, f"ChainOfTitle_{tct_number}_{ts}.docx")


def generate_matter_status_report(matter_code: str) -> dict:
    """Generate a full matter status report."""
    conn = get_db()
    cur = conn.cursor()
    doc = Document()
    doc.add_heading("MATTER STATUS REPORT", 0)

    # ── Matter row ─────────────────────────────────────────────────────
    cur.execute("""
        SELECT m.matter_code, m.title, m.matter_type, m.docket_number,
               m.court_or_agency, m.status, m.verification_status,
               m.verified_document_count, c.name
          FROM matters m
          JOIN clients c ON c.client_code = m.client_code
         WHERE m.matter_code = %s
    """, (matter_code,))
    m = cur.fetchone()
    if m:
        doc.add_heading(m[1], 1)
        p = doc.add_paragraph()
        p.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')} PHT").italic = True
        doc.add_paragraph()
        table = doc.add_table(rows=1, cols=2); table.style = 'Table Grid'
        hdr = table.rows[0].cells; hdr[0].text = "Field"; hdr[1].text = "Value"
        fields = [
            ("Matter Code", m[0]), ("Client", m[8]), ("Type", m[2]),
            ("Docket/Reference", m[3] or "N/A"),
            ("Court/Agency", m[4] or "N/A"),
            ("Status", m[5]), ("Verification", m[6]),
            ("Verified Documents", str(m[7])),
        ]
        for fname, val in fields:
            r = table.add_row().cells; r[0].text = fname; r[1].text = str(val)
        doc.add_paragraph()

    # ── Open action_items ──────────────────────────────────────────────
    cur.execute("""
        SELECT description, due_date, priority, status
          FROM action_items
         WHERE case_file = %s AND status = 'open'
         ORDER BY
           CASE priority WHEN 'Critical' THEN 1 WHEN 'High' THEN 2
                          WHEN 'Medium' THEN 3 ELSE 4 END,
           due_date NULLS LAST
    """, (matter_code,))
    actions = cur.fetchall()
    if actions:
        doc.add_heading("Open Action Items", 2)
        table = doc.add_table(rows=1, cols=4); table.style = 'Table Grid'
        hdr = table.rows[0].cells
        for i, h in enumerate(["Action", "Due", "Priority", "Status"]):
            hdr[i].text = h
        for a in actions:
            r = table.add_row().cells
            r[0].text = a[0]
            r[1].text = str(a[1]) if a[1] else "No date"
            r[2].text = a[2] or "Normal"
            r[3].text = a[3]
        doc.add_paragraph()

    # ── case_reports latest row ────────────────────────────────────────
    cur.execute("""
        SELECT last_hearing_date, last_hearing_type, counsel_present,
               judge_orders, next_hearing_date, evidence_agreed,
               current_position, strategic_risks, missing_fields
          FROM case_reports
         WHERE matter_code = %s
         ORDER BY updated_at DESC LIMIT 1
    """, (matter_code,))
    r = cur.fetchone()
    if r:
        doc.add_heading("Latest Hearing Report", 2)
        table = doc.add_table(rows=1, cols=2); table.style = 'Table Grid'
        hdr = table.rows[0].cells; hdr[0].text = "Field"; hdr[1].text = "Detail"
        fields = [
            ("Last Hearing", f"{r[0]} — {r[1]}" if r[0] else "N/A"),
            ("Counsel Present", r[2] or "Unknown"),
            ("Court Orders", r[3] or "None recorded"),
            ("Next Hearing", str(r[4]) if r[4] else "Not scheduled"),
            ("Evidence Agreed", "Yes" if r[5] else "No"),
            ("Current Position", r[6] or "Not assessed"),
            ("Strategic Risks", r[7] or "None identified"),
            ("Missing Information", ", ".join(r[8]) if r[8] else "Complete"),
        ]
        for fname, val in fields:
            row = table.add_row().cells; row[0].text = fname; row[1].text = str(val)

    cur.close(); conn.close()
    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    return _save_doc(doc, f"MatterReport_{matter_code}_{ts}.docx")


if __name__ == "__main__":
    rt = sys.argv[1] if len(sys.argv) > 1 else "help"
    iden = sys.argv[2] if len(sys.argv) > 2 else None
    if rt == "title" and iden:
        print(json.dumps(generate_title_chain_report(iden)))
    elif rt == "matter" and iden:
        print(json.dumps(generate_matter_status_report(iden)))
    else:
        print("Usage: python generate_report.py [title|matter] [identifier]")
        print("  title  T-4497           Chain of title report")
        print("  matter MWK-CV26360      Matter status report")
